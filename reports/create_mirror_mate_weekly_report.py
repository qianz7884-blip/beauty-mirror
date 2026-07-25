from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Mirror_Mate_weekly_report_product_plan.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(90, 90, 90)
FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"


def set_run_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_para_font(paragraph, size=11, color=INK):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style = doc.styles[f"Heading {level}"]
    p.style = style
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        set_run_font(run, bold=True)
        rest = text[len(bold_lead):]
        if rest:
            run = p.add_run(rest)
            set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [Inches(6.3)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.15
    run = p2.add_run(body)
    set_run_font(run, size=10.5)
    doc.add_paragraph()
    return table


def add_comparison_table(doc):
    headers = ["维度", "美丽修行公开能力", "Mirror Mate 当前状态", "差距判断"]
    rows = [
        [
            "产品定位",
            "面向大众消费者的化妆品查询、成分理解、肤质测评和消费决策平台。",
            "更接近个人美妆/护肤管理工具，强调产品柜、日记、肤况记录和轻量建议。",
            "不适合直接按平台型产品硬对标，应转向个人化记录工具。",
        ],
        [
            "产品与成分数据",
            "公开资料显示已覆盖 350万+ 产品、13万+ 品牌、2.7万+ 成分，并沉淀大量用户心得。",
            "当前以小规模本地产品知识库和静态成分规则为主，适合演示和规则推荐。",
            "最大差距。短期不追求规模，先追求字段质量和推荐可解释。",
        ],
        [
            "AI 测肤",
            "对外提供 AI 测肤、16种肤质分型、8大维度皮肤问题等完整包装。",
            "当前更适合做肤况观察和趋势记录，本地算法可信度不足以作为专业诊断。",
            "应降级为日常参考，不做医学化承诺。",
        ],
        [
            "推荐逻辑",
            "依赖大规模产品库、成分库、肤质模型、用户评价和平台榜单。",
            "可基于用户肤质、已有产品、日记和肤况记录做轻量规则推荐。",
            "推荐可以做，但必须解释依据，不能装成权威判断。",
        ],
        [
            "用户系统",
            "成熟商业 App，具备账号、社区/内容、会员等长期运营基础。",
            "已补齐本地匿名用户管理，可查看当前身份、数据统计、数据库状态并导出备份。",
            "适合演示和小范围试用；正式上线前仍需注册登录和云端备份。",
        ],
        [
            "商业闭环",
            "具备会员、试用盒、榜单、品牌合作、内容运营等商业化空间。",
            "目前还没有明确付费理由，更适合先验证“用户是否愿意连续记录”。",
            "先做留存闭环，再谈盈利。",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_widths(table, [Cm(2.5), Cm(5.0), Cm(5.0), Cm(4.0)])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, size=10.5, color=DARK_BLUE)
        set_cell_shading(table.cell(0, idx), FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            set_cell_text(cells[idx], item, bold=(idx == 0), size=9.5)
    doc.add_paragraph()


def add_plan_table(doc):
    headers = ["阶段", "核心目标", "具体产出", "验收标准"]
    rows = [
        [
            "第 1 周",
            "稳定数据和产品库基础",
            "整理产品字段；补 50-100 个高质量样例；完善本地用户数据、备份和数据库状态提示。",
            "用户能添加产品、看到本地身份和数据统计；推荐可读取规范字段。",
        ],
        [
            "第 2 周",
            "建立个人肤质档案和可解释推荐",
            "新增肤质、敏感程度、主要困扰、避雷成分、偏好功效；重做推荐理由。",
            "每条推荐都能说明：依据的肤质/肤况/产品字段是什么。",
        ],
        [
            "第 3 周",
            "打通日记、肤况和产品使用闭环",
            "日记关联当天使用产品；肤况记录支持趋势展示；弱化单次 AI 准确判断。",
            "用户能看见过去一周使用产品和肤况变化的关系。",
        ],
        [
            "第 4 周",
            "小范围试用和周报验证",
            "找 5-10 名同学试用；记录卡点、留存、功能反馈；形成下一版迭代清单。",
            "至少获得 5 份有效反馈，并明确哪些功能值得继续做。",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_widths(table, [Cm(2.2), Cm(4.1), Cm(5.8), Cm(4.4)])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, size=10.5, color=DARK_BLUE)
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            set_cell_text(cells[idx], item, bold=(idx == 0), size=9.5)
    doc.add_paragraph()


def add_risk_table(doc):
    headers = ["风险", "当前判断", "应对方式"]
    rows = [
        ["数据规模不足", "短期无法追上平台级产品库。", "先做高质量小库，围绕 50-100 个常见产品做完整字段和推荐解释。"],
        ["AI 测肤不够准", "本地算法缺少训练集和验证指标。", "改成“肤况观察/趋势记录”，避免专业诊断表达。"],
        ["用户不愿长期记录", "记录成本是最大留存障碍。", "减少填写字段，优先支持拍照、快速选择和自动带出产品信息。"],
        ["商业化压力过早", "现阶段不具备直接收费基础。", "先用 5-10 人试用验证真实需求，再讨论登录、云端和付费。"],
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [Cm(3.2), Cm(5.2), Cm(7.6)])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, size=10.5, color=DARK_BLUE)
        set_cell_shading(table.cell(0, idx), FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            set_cell_text(cells[idx], item, bold=(idx == 0), size=9.5)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6 if name != "Heading 3" else 4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mirror Mate 项目周报")
    set_run_font(run, size=22, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("与美丽修行的差距分析及下一阶段产品规划")
    set_run_font(run, size=14, bold=True, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"日期：2026年7月19日    项目：Mirror Mate")
    set_run_font(run, size=10.5, color=MUTED)

    add_callout(
        doc,
        "本周核心结论",
        "Mirror Mate 不应以“复制美丽修行”为短期目标。更可行的定位是：个人护肤记录 + 产品柜 + 肤况趋势 + 可解释建议。短期目标应从平台级产品库/专业测肤，调整为能让小范围用户连续使用 2-4 周的记录型工具。",
    )

    add_heading(doc, "一、当前项目状态", 1)
    add_body(
        doc,
        "目前项目已经具备产品管理、护肤/美妆日记、肤况分析展示、产品知识库雏形、推荐逻辑和本地匿名用户管理等基础模块。最新补充的本地用户管理可以区分当前浏览器身份，并在“我的”页面展示当前身份、数据数量、数据库状态和数据导出入口。",
    )
    add_body(
        doc,
        "需要注意的是，当前用户体系还不是正式账号登录。它适合本地演示和小范围试用，但不具备多设备同步、注册登录、云端备份和商业运营所需的账号能力。",
    )
    add_bullet(doc, "已有能力：产品柜、产品添加/识别入口、护肤日记、肤况记录、基础推荐、产品知识库、本地用户数据管理。")
    add_bullet(doc, "主要短板：产品库规模小、成分库深度不足、AI 测肤缺少验证指标、推荐可信度仍需加强。")
    add_bullet(doc, "适合定位：先做个人化记录工具，而不是直接做平台型化妆品数据库。")

    add_heading(doc, "二、与美丽修行的差距对比", 1)
    add_body(
        doc,
        "公开资料显示，美丽修行已经是成熟的化妆品成分与消费决策平台，具备大规模产品库、品牌库、成分库、用户心得和 AI 测肤包装能力。Mirror Mate 当前仍处于 MVP/课程项目阶段，两者不应按同一商业成熟度评价。",
    )
    add_comparison_table(doc)

    add_heading(doc, "三、差距背后的原因", 1)
    add_number(doc, "数据资产不是短期工程。美丽修行的核心壁垒不是页面，而是长期积累的产品、成分、品牌、评价和用户行为数据。")
    add_number(doc, "AI 测肤需要训练数据和验证指标。没有标注数据、测试集和专家校验时，不能把本地视觉算法包装成专业诊断。")
    add_number(doc, "商业闭环需要先有留存。当前 Mirror Mate 还没有足够强的用户持续使用理由，现阶段应先验证记录价值。")
    add_number(doc, "一个人推进时应避免大平台路线。直接复制美丽修行会导致范围失控，反而难以完成可交付成果。")

    add_heading(doc, "四、建议重新确定自己的产品方向", 1)
    add_callout(
        doc,
        "建议定位",
        "Mirror Mate 是面向个人用户的美妆护肤管理工具，帮助用户记录自己的产品、护肤行为和肤况变化，并基于个人记录提供可解释的日常建议。",
    )
    add_body(doc, "这个定位的优势是更适合单人开发，也更容易在周报和答辩中解释其合理性：项目不追求替代大型平台，而是从个人使用场景切入。")
    add_bullet(doc, "产品柜：记录用户已有产品、使用进度、购买/过期时间。")
    add_bullet(doc, "肤质档案：记录肤质、敏感程度、主要困扰、避雷成分和偏好功效。")
    add_bullet(doc, "日记闭环：记录当天使用了哪些产品、肤况如何、是否有不适反应。")
    add_bullet(doc, "可解释推荐：给出“为什么这样建议”，而不是只给一个结论。")
    add_bullet(doc, "趋势复盘：一周或两周后展示肤况变化与使用记录之间的关系。")

    add_heading(doc, "五、下一阶段要完成的产品闭环", 1)
    add_body(doc, "下一阶段不建议继续堆功能，而应围绕一个完整使用流程打磨：")
    add_number(doc, "用户添加自己的产品，系统尽量自动补全品牌、品类、功效、核心成分和注意点。")
    add_number(doc, "用户设置肤质档案，包括肤质、敏感程度、主要困扰和避雷偏好。")
    add_number(doc, "用户每天快速记录肤况和使用产品，可以拍照，也可以手动选择泛红、干燥、爆痘等状态。")
    add_number(doc, "系统根据肤质档案、当天肤况和已有产品，给出轻量、可解释、低风险的今日建议。")
    add_number(doc, "系统按周生成趋势复盘，让用户看到记录带来的价值。")

    add_heading(doc, "六、四周执行计划", 1)
    add_plan_table(doc)

    add_heading(doc, "七、风险与应对", 1)
    add_risk_table(doc)

    add_heading(doc, "八、需要老师确认或支持的事项", 1)
    add_bullet(doc, "确认项目目标是否从“直接盈利的大平台”调整为“先完成可试用 MVP 并验证需求”。")
    add_bullet(doc, "确认是否允许以小范围同学试用作为阶段成果，而不是立即商业化上线。")
    add_bullet(doc, "确认产品库数据是否可以先使用公开信息整理的样例库，并明确数据来源与免责声明。")
    add_bullet(doc, "确认 AI 测肤部分是否可以定位为“日常观察与趋势参考”，避免专业医疗或皮肤诊断承诺。")

    add_heading(doc, "九、可向老师汇报的表述", 1)
    add_body(
        doc,
        "美丽修行属于平台型化妆品成分与消费决策产品，其优势主要来自多年积累的产品库、成分库、用户心得和商业化体系。Mirror Mate 当前不适合直接以该平台为短期对标对象。下一步我计划将项目定位调整为个人化护肤记录工具，重点完成产品柜、肤质档案、日记记录、可解释推荐和周趋势复盘的闭环，并通过小范围试用验证用户是否愿意持续记录。",
    )
    add_body(
        doc,
        "这样可以在资源有限、个人开发的条件下形成可展示、可迭代、可验证的阶段成果，也能为后续是否扩充产品库、接入正式账号系统和迁移云端数据库提供依据。",
    )

    add_heading(doc, "资料来源", 1)
    add_body(doc, "1. 美丽修行官网：https://www.bevol.com/index.html")
    add_body(doc, "2. App Store 美丽修行应用介绍：https://apps.apple.com/cn/app/id1089854728")
    add_body(doc, "3. 小米应用商店美丽修行页面：https://app.mi.com/details?id=cn.bevol.p")
    add_body(doc, "4. Mirror Mate 项目当前代码与本地功能状态，检查日期：2026年7月19日。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    output = build_doc()
    print(output)
