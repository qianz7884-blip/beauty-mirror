"""生成 2026-07-04 ~ 07-05 工作周报 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

# ===================== 封面标题 =====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Beauty Mirror 项目周报')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0x59)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI 美妆与形象管理全栈应用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x8B, 0x7E)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('仓库：github.com/qianz7884-blip/beauty_mirror\n报告日期：2026年7月4日 — 7月5日')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===================== 一、工作概览 =====================
doc.add_heading('一、工作概览', level=1)

doc.add_paragraph(
    '本周（7/4 - 7/5）主要围绕两个方向：一是 Render 云端部署的持续调试与修复，'
    '解决了 Docker 容器中 MediaPipe 依赖缺失和路由异常问题；'
    '二是前端交互优化，为跟镜流程和镜前分析增加了心情选择与满意度反馈闭环。'
)

# 提交记录表
doc.add_heading('Git 提交记录', level=2)
commits = [
    ('2026-07-05', 'f1d2d94', '添加 /api/health 端点 + 构建诊断 + gunicorn --preload', '健康检查 + 启动防护'),
    ('2026-07-05', '11a6e7d', 'COPY 提前到 pip 之前，破坏 Docker 缓存确保新代码部署', 'Docker 缓存修复'),
    ('2026-07-05', 'd26464d', '移除不存在的 mesa-dev 包，只保留运行时库', '包名兼容修复'),
    ('2026-07-05', '82cbff8', '添加 Mesa 渲染后端 + 阻止 pip 安装 GUI 版 OpenCV', 'OpenCV headless 方案'),
    ('2026-07-05', 'a043ec2', '用 headless OpenCV 替代 GUI 版，使用 Mesa 开发库', 'libGLESv2 根因修复'),
    ('2026-07-05', '5e2c535', '添加 DRI 软件渲染驱动 + 强制软件渲染环境变量', 'MediaPipe 兼容'),
    ('2026-07-05', 'e731ae8', '添加 Mesa 渲染后端库，恢复 ldconfig', '系统库配置'),
    ('2026-07-04', 'd93b848', '修复 MediaPipe Linux 系统库缺失', '部署问题初始排查'),
    ('2026-07-04', '2e7b9cf', '修复 Dockerfile：使用 Debian 11 基础镜像 + 正确 OpenGL 包名', '基础镜像切换'),
    ('2026-07-04', 'e5a0e34', '切换到 Docker 部署，修复 MediaPipe 缺少 OpenGL 系统库', '部署方案切换'),
]

table = doc.add_table(rows=len(commits) + 1, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['日期', '提交哈希', '说明', '类别']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

for i, (dt, hsh, msg, cat) in enumerate(commits):
    table.rows[i + 1].cells[0].text = dt
    table.rows[i + 1].cells[1].text = hsh
    table.rows[i + 1].cells[2].text = msg
    table.rows[i + 1].cells[3].text = cat

doc.add_paragraph()

doc.add_page_break()

# ===================== 二、功能开发 =====================
doc.add_heading('二、功能开发', level=1)

# 1. 跟镜流程心情+满意度
doc.add_heading('1. 跟镜流程页 — 心情与满意度反馈', level=2)
doc.add_paragraph(
    '在 /tutorial 跟镜流程页，用户完成妆容步骤后点击"完成，拍妆后照"，'
    '拍照后新增两个交互环节：'
)
items = [
    '今日心情选择：复用 DiaryForm 中的颜色块心情组件（5 种状态：状态极佳/开心/状态稳定/一般/需要慢一点），每个心情有独立的颜色标识',
    '满意度反馈：满意 → 自动创建日记（带心情 + 关联产品 + 完成照）；不满意 → 直接跳过，不保存',
    '未选心情时"满意，保存到日记"按钮禁用，防止漏选',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '前端修改集中在 Tutorial.jsx，引入 MOOD_OPTIONS 心情组件，'
    '状态变量 feedbackMood 贯穿拍照 → 心情 → 保存全流程。'
)

# 2. 镜前分析流程
doc.add_heading('2. 镜前分析页 — 同样增加反馈闭环', level=2)
doc.add_paragraph(
    '在 SkinAnalysisPanel.jsx 的肤质分析结果页底部，同样增加了心情选择和满意度反馈。'
    '用户查看完分析建议后，需要先选择今日心情和满意度才能点击"完成"。'
    '满意时自动创建日记（关联本次分析记录），不满意则直接关闭。'
)

doc.add_page_break()

# ===================== 三、部署调试 =====================
doc.add_heading('三、Render 云部署调试', level=1)

doc.add_paragraph(
    '本周部署调试是耗时最多的工作，经历了 Docker 方案的多轮试错，最终回归原生 Python 运行时。'
)

doc.add_heading('问题链路', level=2)

problems = [
    ('libGLESv2.so.2 缺失',
     'MediaPipe 启动报错 "cannot open shared object file"。'
     '经排查，根因不是 MediaPipe 本身，而是其依赖的 opencv-contrib-python（GUI版）'
     '动态链接了 libGLESv2。解决方案：pip 安装 opencv-contrib-python-headless + '
     'grep -v mediapipe 排重 + --no-deps 安装 mediapipe。'),
    ('Debian 11 包名不兼容',
     '初次尝试安装 mesa-common-dev、libegl1-mesa-dev 等包失败，'
     '这些包在 python:3.11-bullseye (Debian 11) 中不存在。'
     '修正为运行时包：libegl-mesa0、libglx-mesa0、libgl1-mesa-dri。'),
    ('Docker 构建缓存',
     'COPY . . 放在 pip install 之后导致代码变更不触发热层重建。'
     '将 COPY 提前到 pip install 之前，利用代码变更破坏缓存。'),
    ('路由 404（未解决）',
     '切换到 Docker 运行时后，构建成功、服务 Live，但所有 Flask 路由返回 404。'
     '尝试了构建时预检、移除 --preload、简化 Dockerfile 等方式均无效。'
     '最终放弃 Docker 方案，回归 Render 原生 Python 运行时。'),
    ('健康检查回滚',
     '原 healthCheckPath 指向 /api/dashboard，当服务异常时 Render 自动回滚到旧部署。'
     '添加零依赖的 /api/health 端点，修改 render.yaml 健康检查路径。'),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('技术总结', level=2)

table2 = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['问题', '根因', '最终方案']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

summary_data = [
    ('libGLESv2.so.2 缺失', 'opencv-contrib-python (GUI) 动态链接 GL', 'pip install opencv-contrib-python-headless'),
    ('mesa-dev 包不存在', 'Debian 11 bullseye 无此包名', '改用 libegl-mesa0 等运行时包'),
    ('Docker 路由 404', '未定位（疑似 Flask+gunicorn 兼容问题）', '回归 Render 原生 Python 运行时'),
    ('自动回滚到旧部署', '健康检查 /api/dashboard 失败', '添加 /api/health 端点 + 修改 render.yaml'),
]
for i, (problem, cause, fix) in enumerate(summary_data):
    table2.rows[i + 1].cells[0].text = problem
    table2.rows[i + 1].cells[1].text = cause
    table2.rows[i + 1].cells[2].text = fix

doc.add_page_break()

# ===================== 四、前端交互优化 =====================
doc.add_heading('四、前端交互优化详情', level=1)

doc.add_heading('1. 跟镜流程页改动（Tutorial.jsx）', level=2)
doc.add_paragraph(
    'handleSaveDiary() 中 mood 由硬编码 "stable" 改为用户选择的 feedbackMood。'
    '切换时间/场景时自动重置心情状态。'
    '"满意"按钮新增 !feedbackMood 禁用条件，确保用户在保存前完成心情选择。'
)

doc.add_heading('2. 镜前分析页改动（SkinAnalysisPanel.jsx）', level=2)
doc.add_paragraph(
    '在分析结果与操作按钮之间插入心情选择区（MOOD_OPTIONS 颜色块）和满意度选择区。'
    '"完成"按钮逻辑：满意时先调用 createDiary 自动保存日记，'
    '再调用 onClose 关闭面板；不满意直接关闭。'
    'handleRetry 中重置 feedbackMood 和 feedbackSatisfaction。'
)

doc.add_heading('3. 心情组件复用', level=2)
doc.add_paragraph(
    '心情颜色块组件最早出现在 DiaryForm.jsx，定义于 utils/moods.js。'
    'MOOD_OPTIONS 包含 5 种心情，每个有独立的 key、label、color、bg。'
    'CSS 类名 dv-mood-selector / dv-mood-option / dv-mood-swatch 定义于 diary.css，'
    '全局可用，无需额外样式修改。'
)

doc.add_page_break()

# ===================== 五、文件变更统计 =====================
doc.add_heading('五、文件变更统计', level=1)

file_changes = [
    ('backend/Dockerfile', '多轮修改，最终删除', 'Docker 方案试错 → 回归原生 Python'),
    ('render.yaml', '修改健康检查路径 + 环境变量', 'healthCheckPath: /api/health'),
    ('backend/routes/health.py', '新建', '零依赖健康检查端点'),
    ('backend/routes/__init__.py', '修改', '注册 health_bp'),
    ('backend/app.py', '修改（已还原）', '添加 fallback 错误处理'),
    ('frontend/src/pages/Tutorial.jsx', '修改', '添加心情选择 + 满意度反馈'),
    ('frontend/src/components/SkinAnalysisPanel.jsx', '修改', '添加心情选择 + 满意度反馈'),
]

table3 = doc.add_table(rows=len(file_changes) + 1, cols=3, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['文件', '操作', '说明']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (f, op, desc) in enumerate(file_changes):
    table3.rows[i + 1].cells[0].text = f
    table3.rows[i + 1].cells[1].text = op
    table3.rows[i + 1].cells[2].text = desc

doc.add_paragraph()

# ===================== 六、后续规划 =====================
doc.add_heading('六、后续规划', level=1)

pending = [
    ('MediaPipe 云部署兼容', '在 Render 原生 Python 运行时中解决 opencv/libGLESv2 依赖，恢复镜前肤质分析功能'),
    ('心情反馈持久化', '满意/不满意反馈数据同步到后端数据库，支持历史查看和统计分析'),
    ('心情组件抽象', '将 dv-mood-selector 抽象为独立 React 组件，避免 DiaryForm/Tutorial/SkinAnalysisPanel 三处重复'),
    ('流程页 AI 拍照指导', '在跟镜流程拍照环节接入 AI，自动检测妆容完成度并给出建议'),
]

table4 = doc.add_table(rows=len(pending) + 1, cols=2, style='Light Shading Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['计划', '说明']):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (feat, desc) in enumerate(pending):
    table4.rows[i + 1].cells[0].text = feat
    table4.rows[i + 1].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph()

# 页脚
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Beauty Mirror · AI 美妆与形象管理助手 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# 保存
output_path = 'f:/学习周报/beauty_mirror/reports/BeautyMirror_研究生周报_2026-07-05_v2.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
