"""生成工作周报 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)

# ===================== 封面标题 =====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Beauty Mirror 项目工作周报')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0x59)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI 美妆与形象管理全栈应用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x8B, 0x7E)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'仓库：github.com/qianz7884-blip/beauty_mirror\n报告日期：{date.today()}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===================== 一、工作概览 =====================
doc.add_heading('一、工作概览', level=1)

doc.add_paragraph(
    '2026年6月期间，独立完成了 Beauty Mirror（AI美妆与形象管理助手）全栈项目从零到一的搭建。'
    '项目采用 React + Vite 前端 + Flask 后端架构，集成了 Google Gemini 视觉识别、'
    'MediaPipe 面部地标检测和 ChromaDB RAG 皮肤科知识库三大 AI 能力，'
    '实现了护肤品库存管理、妆容日记记录、AI拍照识别产品、AI面部肤质分析等核心功能。'
)

# 提交记录表
doc.add_heading('Git 提交记录', level=2)
table = doc.add_table(rows=8, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['日期', '提交哈希', '说明', '影响规模']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

commits = [
    ('2026-06-20', '0fc8fd3', '新增肤质分析、面部区域检测、个人资料页等功能模块', '+4870/-298 行'),
    ('2026-06-14', 'b634903', '粉色和绿色主题样式调整', '主题切换'),
    ('2026-06-14', '78e1435', '准备 Render + Vercel 部署配置', '+render.yaml'),
    ('2026-06-11', '44c25df', '增加 README 项目文档', '+293行'),
    ('2026-06-11', '6a89128', '重构项目结构并同步前端代码', '项目重组'),
    ('2026-06-11', '45c918e', '使用 PORT 环境变量适配云部署', '配置'),
    ('2026-06-11', 'd8fb31a', '首次提交：Beauty Mirror Flask Web 应用', '初始版本'),
]
for i, (dt, hsh, msg, scale) in enumerate(commits):
    table.rows[i+1].cells[0].text = dt
    table.rows[i+1].cells[1].text = hsh
    table.rows[i+1].cells[2].text = msg
    table.rows[i+1].cells[3].text = scale

doc.add_paragraph()

# 工作量统计
doc.add_heading('工作量统计', level=2)

stats = [
    ('后端 Python 代码', '~3,200 行', 'app.py / models.py / config.py / recognizer.py / skin_analyzer.py / face_regions.py / vector_store.py'),
    ('前端 React 代码', '~5,500 行', 'App.jsx / api.js / 5 pages / 9 components / index.css'),
    ('CSS 样式', '3,445 行', '全局样式（莫兰迪鼠尾草绿主题 + iOS毛玻璃 + 渐变背景）'),
    ('配置文件', '8 个', 'requirements.txt / package.json / vite.config.js / vercel.json / render.yaml / Procfile / .env / .gitignore'),
    ('文档', '~400 行', 'README.md（294行）+ WEEKLY_REPORT.md（118行）'),
    ('数据模型', '3 个', 'Product / Diary / SkinAnalysis'),
    ('API 接口', '14 个', '产品CRUD + 日记CRUD + AI识别 + 肤质分析(4) + 统计 + 静态资源'),
    ('AI 模型文件', '3.7 MB', 'MediaPipe Face Landmarker 模型（478点面部检测）'),
    ('知识库', '13 条', '皮肤科种子知识（ChromaDB 向量化）'),
]

table2 = doc.add_table(rows=len(stats)+1, cols=3, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['类别', '规模', '详情']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

for i, (cat, cnt, detail) in enumerate(stats):
    table2.rows[i+1].cells[0].text = cat
    table2.rows[i+1].cells[1].text = cnt
    table2.rows[i+1].cells[2].text = detail

doc.add_page_break()

# ===================== 二、功能模块详解 =====================
doc.add_heading('二、功能模块详解', level=1)

# 1. 护肤品管理
doc.add_heading('1. 护肤品管理', level=2)
doc.add_paragraph(
    '完整的 CRUD 功能，支持新增、编辑、删除、搜索和分类筛选。'
    '列表/网格双视图切换，自定义分类管理持久化到 localStorage。'
    '图片上传支持本地文件和 URL 粘贴两种方式，自动生成缩略图。'
)

# 2. 妆容日记
doc.add_heading('2. 妆容日记', level=2)
doc.add_paragraph(
    '记录每日妆容心得，内置心情选择器（😍/😊/😐/😢）。'
    '支持关联多个已添加的护肤品、图片上传。'
    '新增日记详情页（DiaryDetail.jsx），可查看完整内容。'
)

# 3. AI 产品识别
doc.add_heading('3. AI 产品识别', level=2)
doc.add_paragraph(
    '核心亮点功能。拍照或从相册选图后，通过 Google Gemini Vision API 自动识别品牌、产品名称和分类。'
    '识别结果支持手动编辑修正，可一键搜索百度图片查找官方宣传图，确认后直接创建产品入库。'
    '该模块设计为可插拔架构（recognizer.py），后续切换 AI 模型只需改一个文件。'
)

# 4. AI 肤质分析（最重要）
doc.add_heading('4. AI 肤质分析（核心功能）', level=2)
doc.add_paragraph(
    '这是本轮开发中工作量最大的模块，实现了完整的 AI 分析管线。'
    '用户拍照上传面部照片后，系统依次执行：'
)

analysis_steps = [
    ('MediaPipe 面部检测', '使用 Face Landmarker 模型检测 478 个面部特征点，精确定位五官轮廓。'),
    ('人脸 ROI 裁剪', '根据特征点自动定位并裁剪面部区域，去除背景干扰。'),
    ('8 区域分区提取', '将面部分为：前额 / 左脸颊 / 右脸颊 / 鼻子 / 下巴 / 左眼周 / 右眼周 / 唇周，共 8 个区域分别提取图片。'),
    ('RAG 知识增强', '从 ChromaDB 向量库中检索皮肤科专业知识（13条种子知识 + Gemini Embedding 768维向量），为 AI 分析提供专业上下文。'),
    ('Gemini Vision 多图分析', '全脸图 + 8 个分区图共 9 张图片并发送给 Gemini Vision，进行综合评估。'),
    ('热力图生成', '使用 matplotlib + numpy + scipy 生成 Apple Health 风格的面部热力图，直观展示各区域评分分布。'),
    ('结构化输出', '返回肤质类型、全脸5项维度评分（水润度/光滑度/光泽度/毛孔/均匀度）、8区域分区评分、主要问题标签、AI评估总结、编号护理建议。'),
]

for i, (title, desc) in enumerate(analysis_steps):
    p = doc.add_paragraph()
    run = p.add_run(f'步骤 {i+1}：{title}')
    run.bold = True
    p.add_run(f' — {desc}')

doc.add_paragraph()
doc.add_paragraph(
    '分析历史自动保存到数据库，支持回顾和趋势对比。'
    '对应的前端组件 SkinAnalysisPanel.jsx（621行）实现了完整的交互流程：'
    '拍照 → 上传 → 分析进度 → 结果展示（评分卡片 + 热力图 + 护理建议）。'
)

# 5. 首页 Dashboard
doc.add_heading('5. 首页 Dashboard', level=2)
doc.add_paragraph(
    '聚合统计卡片展示产品总数、本月新增、日记数、分析次数。'
    '快速入口区域提供拍照识别、相册识别、手动录入、肤质分析四个快捷操作。'
    '底部展示最近产品和近期肤质分析历史，点击可跳转详情。'
)

# 6. 个人设置
doc.add_heading('6. 个人设置页', level=2)
doc.add_paragraph(
    '新增 Profile.jsx（261行），支持肤质偏好设置（干性/油性/混合/敏感）和每日提醒开关。'
    '界面采用 iOS 风格设置列表，与整体莫兰迪风格统一。'
)

doc.add_page_break()

# ===================== 三、技术架构 =====================
doc.add_heading('三、技术架构', level=1)

doc.add_heading('技术栈', level=2)
tech = [
    ('前端框架', 'React 18 + Vite 5', 'SPA 单页应用，HMR 热更新'),
    ('路由', 'React Router v6', '4 页面 + 404 兜底 + 日记详情'),
    ('HTTP', 'Axios', 'API 调用 + 图片地址工具函数'),
    ('UI 图标', 'Lucide React', '现代化线性图标库'),
    ('后端框架', 'Flask 3', 'RESTful API，14 个接口'),
    ('ORM', 'Flask-SQLAlchemy', 'Product / Diary / SkinAnalysis 三模型'),
    ('数据库', 'SQLite（默认）', '零配置，支持 MySQL/PostgreSQL 切换'),
    ('AI 视觉', 'Google Gemini Vision', '产品识别 + 肤质分析'),
    ('面部检测', 'MediaPipe', 'Face Landmarker 478 点检测'),
    ('向量数据库', 'ChromaDB', 'RAG 皮肤科知识库'),
    ('Embedding', 'Gemini text-embedding-004', '768 维语义向量'),
    ('图像处理', 'Pillow', '缩略图 + 格式转换'),
    ('热力图', 'matplotlib + numpy + scipy', 'Apple Health 风格面部热力图'),
    ('部署', 'Render + Vercel', '后端/前端分离部署，已配置完成'),
]

table3 = doc.add_table(rows=len(tech)+1, cols=3, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['层级', '技术', '说明']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (layer, t, desc) in enumerate(tech):
    table3.rows[i+1].cells[0].text = layer
    table3.rows[i+1].cells[1].text = t
    table3.rows[i+1].cells[2].text = desc

doc.add_paragraph()

doc.add_heading('项目结构', level=2)
doc.add_paragraph(
    '后端（backend/）：app.py（Flask入口+14个API路由）、models.py（3个数据模型）、'
    'config.py（数据库配置）、recognizer.py（AI产品识别）、skin_analyzer.py（AI肤质分析管线689行）、'
    'face_regions.py（MediaPipe 478点→8区域ROI提取）、knowledge_base/（ChromaDB向量库+RAG）'
)
doc.add_paragraph(
    '前端（frontend/）：App.jsx（路由定义）、api.js（Axios封装）、'
    '5个页面（Dashboard/ProductManage/MakeupDiary/DiaryDetail/Profile）、'
    '9个组件（Layout/ProductForm/ProductCard/DiaryForm/DiaryCard/'
    'RecognizePanel/SkinAnalysisPanel/ImageViewer）'
)

doc.add_heading('API 接口一览', level=2)
apis = [
    ('GET', '/api/dashboard', '首页统计数据'),
    ('GET', '/api/products', '产品列表（支持 search & category 筛选）'),
    ('POST', '/api/products', '新增产品（multipart 表单）'),
    ('GET/PUT/DELETE', '/api/products/<id>', '产品详情/编辑/删除'),
    ('GET', '/api/diary', '日记列表'),
    ('POST', '/api/diary', '新增日记'),
    ('GET/PUT/DELETE', '/api/diary/<id>', '日记详情/编辑/删除'),
    ('POST', '/api/recognize', 'AI 产品识别（Gemini Vision）'),
    ('POST', '/api/skin-analysis', 'AI 肤质分析（完整管线）'),
    ('GET', '/api/skin-analyses', '分析历史列表'),
    ('GET/DELETE', '/api/skin-analyses/<id>', '分析详情/删除'),
]

table4 = doc.add_table(rows=len(apis)+1, cols=3, style='Light Shading Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['方法', '路径', '说明']):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (method, path, desc) in enumerate(apis):
    table4.rows[i+1].cells[0].text = method
    table4.rows[i+1].cells[1].text = path
    table4.rows[i+1].cells[2].text = desc

doc.add_page_break()

# ===================== 四、UI/UX设计 =====================
doc.add_heading('四、UI/UX 设计', level=1)
doc.add_paragraph(
    '整体采用移动端优先的卡片式界面设计，主题色从最初粉色演变为莫兰迪鼠尾草绿（粉绿配色）。'
    '实现了以下视觉特性：'
)

ux_items = [
    '底部固定 Tab 导航（首页/产品/日记/我的），图标 + 文字标签，当前页高亮',
    'iOS 风格毛玻璃效果（backdrop-filter: blur），卡片半透明叠加',
    '丝绸/水波纹叠加渐变背景（animation），多层次视觉效果',
    'CSS 变量体系，统一管理主题色、圆角、阴影、间距',
    '响应式设计，适配手机/平板/桌面多终端',
    '产品列表支持列表/网格双视图切换',
    '图片全屏查看器（ImageViewer），支持滑动浏览',
    '心情选择器使用 Emoji 表情符号，直观友好',
    '表单弹层模式，减少页面跳转',
    '肤质分析进度动画，提升等待体验',
]
for item in ux_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===================== 五、部署与环境 =====================
doc.add_heading('五、部署与环境', level=1)
doc.add_paragraph(
    '已完成 Render + Vercel 分离部署配置：'
)

deploy_items = [
    '后端（Render）：render.yaml + Procfile，gunicorn 启动，自动读取 PORT 环境变量',
    '前端（Vercel）：vercel.json SPA 配置，通过 VITE_API_BASE_URL 连接后端',
    '数据库：默认 SQLite 零配置，支持 DATABASE_URL 环境变量切换 MySQL/PostgreSQL',
    '环境变量：GEMINI_API_KEY（AI 核心）、SECRET_KEY、DATABASE_URL 等',
    'AI 模型可配置：GEMINI_MODEL（默认 gemini-2.5-flash）、GEMINI_TIMEOUT（默认 60s）',
    '本地开发：后端 5000 端口，前端 3000 端口（Vite 代理 /api 到后端）',
    '.gitignore 已配置排除 node_modules、dist、__pycache__、instance、上传图片',
]
for item in deploy_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===================== 六、设计亮点 =====================
doc.add_heading('六、技术亮点与设计决策', level=1)

highlights = [
    ('可插拔 AI 架构',
     'AI 识别和分析模块独立封装，recognizer.py 和 skin_analyzer.py 各自内聚，'
     '更换 AI 模型（如从 Gemini 切换到 Claude Vision）只需修改单个文件。'),
    ('完整分析管线',
     '肤质分析不是简单的单图识别，而是 7 步管线：MediaPipe 检测 → ROI 裁剪 → 8 区域提取 → '
     'RAG 知识检索 → Gemini 多图并发 → 热力图生成 → 结构化输出。每一步都可独立调试和优化。'),
    ('RAG 知识增强',
     '13 条皮肤科种子知识经 Gemini Embedding 向量化存入 ChromaDB，'
     '分析时自动检索相关专业知识注入 Prompt，提升 AI 分析的专业性和准确性。'),
    ('面部热力图',
     '使用 matplotlib + numpy + scipy 根据 8 区域评分生成 Apple Health 风格的热力图，'
     '直观展示各区域皮肤状态，是连接数据与用户的可视化桥梁。'),
    ('移动端优先',
     '从 CSS 变量、rem 单位、触控友好的交互到毛玻璃 Tab 导航，'
     '整个 UI 体系围绕手机使用场景设计，也能自适应大屏。'),
    ('零配置数据库',
     '开发环境使用 SQLite，无需安装任何数据库服务。'
     '生产环境只需设置一个 DATABASE_URL 即可无缝切换到 MySQL/PostgreSQL。'),
    ('分离部署',
     '前后端独立部署到不同平台（Vercel + Render），通过环境变量连接，'
     '充分利用各自平台的优势（Vercel 的 CDN + Render 的持久化存储）。'),
]

for i, (title, desc) in enumerate(highlights):
    p = doc.add_paragraph()
    run = p.add_run(f'{i+1}. {title}')
    run.bold = True
    doc.add_paragraph(desc)

doc.add_page_break()

# ===================== 七、后续规划 =====================
doc.add_heading('七、后续规划', level=1)

pending = [
    ('多用户系统', '添加用户注册/登录，隔离个人数据，支持云端备份同步。'),
    ('产品过期提醒', '录入产品保质期，到期前自动推送提醒，减少浪费。'),
    ('肤质趋势分析', '基于历史分析记录生成肤质变化趋势图，追踪护肤效果。'),
    ('社区分享', '妆容日记分享功能，允许用户发布公开日记并互相点赞评论。'),
    ('多语言支持', '当前面向中文用户，后续可扩展英文/日文等多语言。'),
    ('PWA 离线支持', '通过 Service Worker 实现离线访问，提升移动端体验。'),
    ('CI/CD', '接入 GitHub Actions，自动测试 + 自动部署。'),
]

table5 = doc.add_table(rows=len(pending)+1, cols=2, style='Light Shading Accent 1')
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['功能', '说明']):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (feat, desc) in enumerate(pending):
    table5.rows[i+1].cells[0].text = feat
    table5.rows[i+1].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph()

# 页脚
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Beauty Mirror · AI 美妆与形象管理助手 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# 保存
output_path = 'f:/学习周报/beauty_mirror/Beauty_Mirror_工作周报.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
